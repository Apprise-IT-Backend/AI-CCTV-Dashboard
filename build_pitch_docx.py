"""Build PLATFORM_PITCH.docx — 4-page pitch doc.

Sections:
  1. Our platform vs Milestone (advantages-first)
  2. Local install hardware
  3. Cloud requirements
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BLUE = RGBColor(0x1F, 0x4E, 0x78)
GREY = RGBColor(0x55, 0x55, 0x55)


def shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def style_header_row(row):
    for cell in row.cells:
        shade(cell, '1F4E78')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)


def add_heading(doc, text, level, color=BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.size = Pt(10)
        r.font.bold = True
        r = p.add_run(' — ' + text)
        r.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None, zebra=True, small=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    style_header_row(t.rows[0])
    body_size = Pt(9) if small else Pt(10)
    for r_idx, row_data in enumerate(rows):
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            if zebra and r_idx % 2 == 1:
                shade(row.cells[i], 'F2F6FA')
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = body_size
                    if i == 0:
                        run.font.bold = True
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    return t


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ══════════════════════════════════════════════════════════════
    # PAGE 1 — Title + Why choose us
    # ══════════════════════════════════════════════════════════════
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('AI-CCTV Dashboard')
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = BLUE

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Why Choose Our Platform Over Milestone XProtect')
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = GREY

    add_heading(doc, 'The Short Answer', level=1)
    add_para(doc,
        'Milestone XProtect is a huge, recording-first system built for large enterprises. '
        'It is powerful, but it is also expensive, complex, and needs paid add-ons for the '
        'features people actually want today — like face recognition and fire detection.'
    )
    add_para(doc,
        'Our platform gives you those AI features built-in, on day one, with no license '
        'fees. It is faster to set up, cheaper to run, and easier to use.', bold=True
    )

    add_heading(doc, 'Our Top Advantages', level=1)

    add_bullet(doc, 'The AI features you actually want are already inside — face '
                    'recognition, person detection, and fire/smoke alerts. XProtect makes '
                    'you buy separate add-ons (like BriefCam) for these.',
               bold_prefix='Built-in AI, no add-ons')

    add_bullet(doc, 'No per-camera license fees. No add-on modules to buy. Open source, '
                    'you own the code and the data.',
               bold_prefix='Zero licensing cost')

    add_bullet(doc, 'One command (start.bat) launches the whole stack. A new user can be '
                    'watching live camera feeds in under 5 minutes.',
               bold_prefix='Instant setup')

    add_bullet(doc, 'Upload a photo of a person and the system starts recognizing them '
                    'within seconds — no restart, no retraining wait.',
               bold_prefix='Hot-reload face enrollment')

    add_bullet(doc, 'Every recognized person can be tagged as VIP, Staff, Visitor, or '
                    'Threat, with notes attached. XProtect has no direct equivalent.',
               bold_prefix='Smart person categories')

    add_bullet(doc, 'Click a person\'s name and see their movement history on a map — '
                    'which cameras saw them, when, and where they went.',
               bold_prefix='Movement path tracking')

    add_bullet(doc, 'Each user gets their own cameras, their own alerts, and their own '
                    'settings. Perfect for shared installations.',
               bold_prefix='True multi-user separation')

    add_bullet(doc, 'Fire detection uses a real AI model, and falls back to a smart color '
                    'and motion check if the model is missing. XProtect relies on '
                    'third-party plugins.',
               bold_prefix='Fire & smoke detection')

    add_bullet(doc, 'Runs on a single modest server. XProtect is architected for a farm '
                    'of servers with a much bigger footprint.',
               bold_prefix='Lightweight footprint')

    add_bullet(doc, 'Simple browser dashboard, no software to install on viewers. '
                    'Works on any modern laptop or tablet.',
               bold_prefix='Zero-install viewing')

    # ══════════════════════════════════════════════════════════════
    # PAGE 2 — Side by side comparison
    # ══════════════════════════════════════════════════════════════
    page_break(doc)

    add_heading(doc, 'Side-by-Side: Where We Win', level=1)
    add_para(doc,
        'A direct comparison on the things that matter most to a small or mid-size site.'
    )

    add_table(doc,
        ['What You Need', 'Our Platform', 'Milestone XProtect'],
        [
            ['Face recognition',
             'Built-in and free',
             'Needs BriefCam paid add-on'],
            ['Fire & smoke alerts',
             'Built-in and free',
             'Third-party plugin required'],
            ['Person detection',
             'Built-in (YOLO AI)',
             'Add-on required'],
            ['Setup time',
             'Under 5 minutes',
             'Days, needs training'],
            ['License fee',
             'None — open source',
             'Paid per camera + modules'],
            ['Add a new person',
             'Upload photo, live in seconds',
             'Separate BriefCam enrollment'],
            ['Person tagging (VIP/Staff/Threat)',
             'Yes, with notes',
             'Not built-in'],
            ['Movement history on map',
             'Yes, click and see path',
             'Not built-in'],
            ['Per-user data separation',
             'Yes, out of the box',
             'Requires role configuration'],
            ['Server hardware',
             'One modest PC',
             'Multi-server farm typical'],
            ['Viewer software',
             'Any web browser',
             'Smart Client desktop app'],
            ['Maintenance',
             'Restart and go',
             'Contracts and support tiers'],
        ],
        col_widths=[Inches(2.1), Inches(2.5), Inches(2.5)],
        small=True
    )

    add_heading(doc, 'Where Milestone Still Leads (Fair Disclosure)', level=2)
    add_para(doc,
        'XProtect is stronger if you need 24/7 continuous recording for weeks or months, '
        'court-grade signed evidence, single sign-on with Active Directory, thousands of '
        'cameras across many buildings, or a native mobile app for security guards. If any '
        'of those are non-negotiable, XProtect is the safer pick.'
    )
    add_para(doc,
        'For everyone else — small businesses, schools, hotels, offices, shops, warehouses '
        '— our platform delivers 90% of what people actually use, at a fraction of the cost '
        'and complexity.', italic=True
    )

    # ══════════════════════════════════════════════════════════════
    # PAGE 3 — Local hardware
    # ══════════════════════════════════════════════════════════════
    page_break(doc)

    add_heading(doc, 'Local Install: Hardware Requirements', level=1)
    add_para(doc,
        'Sizing below is for a 10-camera install running all AI features (face, person, '
        'fire) in real time. The AI engine is the main resource user; everything else is '
        'lightweight.'
    )

    add_heading(doc, 'Three Tiers to Choose From', level=2)

    add_table(doc,
        ['', 'Minimum', 'Recommended', 'Comfortable'],
        [
            ['Purpose',
             'Runs, but tight',
             'The sweet spot',
             'Room to grow'],
            ['CPU',
             '8-core (i7-10700 / Ryzen 7 3700X)',
             '6-core (i5-12400 / Ryzen 5 5600)',
             '8-12 core (i7-13700 / Ryzen 7 7700)'],
            ['RAM',
             '16 GB',
             '16 GB',
             '32 GB'],
            ['Storage',
             '256 GB SSD',
             '512 GB NVMe',
             '1 TB NVMe'],
            ['GPU',
             'None (CPU only)',
             'NVIDIA GTX 1650 / RTX 3050 (4 GB+)',
             'NVIDIA RTX 4060 / 4070 (8-12 GB)'],
            ['Network',
             'Wired Gigabit',
             'Wired Gigabit',
             'Wired 2.5 GbE'],
            ['Trade-off',
             'Drop to 3-4 FPS, turn off person detection',
             'All AI features on, smooth',
             'Handles up to 20-30 cameras'],
            ['Est. cost',
             '$500-700',
             '$700-900',
             '$1,200-1,600'],
        ],
        col_widths=[Inches(1.2), Inches(2.0), Inches(2.0), Inches(2.0)],
        small=True
    )

    add_heading(doc, 'What Uses the Resources', level=2)
    add_bullet(doc,
        'The AI engine (face + person + fire detection) is 70% of the workload. '
        'A basic NVIDIA GPU roughly cuts CPU use in half and makes everything smooth.',
        bold_prefix='AI processing')
    add_bullet(doc,
        '10 cameras at 1080p use 20-40 Mbps of your LAN. Wired connections only — '
        'WiFi will drop frames.',
        bold_prefix='Network')
    add_bullet(doc,
        'Only snapshots are saved (not full video), so 500 GB SSD is plenty for '
        'months of history.',
        bold_prefix='Storage')
    add_bullet(doc,
        'Cameras must be reachable on the same LAN as the server. Use a wired switch, '
        'not the wireless router.',
        bold_prefix='Cameras')

    add_heading(doc, 'Our Honest Recommendation', level=2)
    add_para(doc,
        'For a real 10-camera install, get the Recommended tier — a mid-range desktop '
        'with any modern NVIDIA graphics card. Around $800 total, it will handle all '
        '10 cameras with every AI feature switched on and still have headroom.',
        bold=True
    )

    # ══════════════════════════════════════════════════════════════
    # PAGE 4 — Cloud
    # ══════════════════════════════════════════════════════════════
    page_break(doc)

    add_heading(doc, 'Cloud Deployment: What You Need', level=1)
    add_para(doc,
        'Instead of buying a local server, you can rent one in the cloud (AWS, Azure, '
        'or Google Cloud). The software works the same way — you just point the cameras '
        'at the cloud address.'
    )

    add_heading(doc, 'The Biggest Thing to Know: Your Internet Upload Speed', level=2)
    add_para(doc,
        '10 cameras need to upload roughly 20-40 Mbps of video 24/7 to the cloud. Most '
        'business fiber lines handle this easily. Home internet often does not. If your '
        'upload speed is under 30 Mbps, use the cameras\' low-resolution "sub-stream" — '
        'that drops the total to about 5-10 Mbps.',
    )

    add_heading(doc, 'Cloud Server Sizing (for 10 Cameras)', level=2)
    add_table(doc,
        ['Provider', 'Instance Type', 'GPU', 'CPU / RAM', 'Monthly Cost'],
        [
            ['AWS',      'g4dn.xlarge',        'NVIDIA T4',    '4 / 16 GB',  '~$380'],
            ['Azure',    'NC4as T4 v3',        'NVIDIA T4',    '4 / 28 GB',  '~$390'],
            ['Google',   'n1-standard-4 + T4', 'NVIDIA T4',    '4 / 15 GB',  '~$350'],
            ['Budget',   'Runpod / Vast.ai',   'RTX 4090',     '8 / 32 GB',  '~$300 (spot ~$150)'],
        ],
        col_widths=[Inches(1.0), Inches(1.8), Inches(1.3), Inches(1.3), Inches(1.6)],
        small=True
    )

    add_heading(doc, 'Full Monthly Cost Estimate — 10 Cameras', level=2)
    add_table(doc,
        ['What You Pay For', 'Cheap', 'Mid', 'Comfortable'],
        [
            ['Cloud server (GPU)',     '$150 (spot)', '$380 (T4)',  '$700 (A10G)'],
            ['Database (managed MySQL)', '$0 (same server)', '$40', '$80'],
            ['Storage for snapshots',  '$5',          '$15',        '$30'],
            ['Bandwidth (1 viewer watching all day)', '$50', '$150', '$300'],
            ['Total per month',        '~$200',       '~$580',      '~$1,100'],
        ],
        col_widths=[Inches(2.4), Inches(1.5), Inches(1.5), Inches(1.5)],
        small=True
    )
    add_para(doc,
        'Note: Uploading to the cloud is usually free. Watching the feeds back on your '
        'phone or laptop is what costs bandwidth. Keep viewers on your local network '
        'where possible.', italic=True, size=9
    )

    add_heading(doc, 'When Cloud Makes Sense', level=2)
    add_bullet(doc, 'You have several sites and want one dashboard for all of them.')
    add_bullet(doc, 'You do not want to buy or maintain a physical server.')
    add_bullet(doc, 'You have fast, reliable internet at each camera location.')
    add_bullet(doc, 'You expect to grow past 10-20 cameras soon.')

    add_heading(doc, 'When Local Wins', level=2)
    add_bullet(doc, 'You have limited internet upload speed.')
    add_bullet(doc, 'You want a one-time cost, not a monthly bill.')
    add_bullet(doc, 'You need the fastest possible response with no internet delay.')
    add_bullet(doc, 'You want camera footage to stay inside your building.')

    add_heading(doc, 'Bottom Line', level=1)
    add_para(doc,
        'For 10 cameras, a local install at around $800 one-time is the cheapest and '
        'simplest choice. Move to cloud when you need to serve multiple sites, when you '
        'do not want to manage hardware, or when you grow past 30 cameras.',
        bold=True
    )

    out = 'PLATFORM_PITCH.docx'
    doc.save(out)
    print(f'Saved {out}')


if __name__ == '__main__':
    main()
