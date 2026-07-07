"""Build CLOUD_AND_FLOW.docx — comfortable cloud cost + simple system diagram."""
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from matplotlib.patches import Rectangle
import os

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


BLUE   = RGBColor(0x1F, 0x4E, 0x78)
GREY   = RGBColor(0x55, 0x55, 0x55)

# Diagram palette
C_USER   = '#2E86AB'
C_CAM    = '#6C757D'
C_MTX    = '#F18F01'
C_AI     = '#C73E1D'
C_BROWSER = '#20A39E'
C_ARROW  = '#333333'
C_BG     = '#F7F9FC'


# ── Diagram generator ──────────────────────────────────────────
def make_diagram(path):
    fig, ax = plt.subplots(figsize=(12, 7.5), dpi=170)
    ax.set_xlim(0, 120)
    ax.set_ylim(0, 80)
    ax.axis('off')
    fig.patch.set_facecolor(C_BG)

    def box(x, y, w, h, text, color, text_color='white', fontsize=12, bold=True):
        b = FancyBboxPatch((x, y), w, h,
                           boxstyle="round,pad=0.6,rounding_size=1.2",
                           linewidth=2, edgecolor=color, facecolor=color)
        ax.add_patch(b)
        ax.text(x + w/2, y + h/2, text,
                ha='center', va='center',
                fontsize=fontsize, color=text_color,
                fontweight='bold' if bold else 'normal',
                wrap=True)

    def arrow(x1, y1, x2, y2, color=C_ARROW, style='-|>', lw=2.2):
        a = FancyArrowPatch((x1, y1), (x2, y2),
                            arrowstyle=style, mutation_scale=22,
                            linewidth=lw, color=color)
        ax.add_patch(a)

    def label(x, y, text, color='#333'):
        ax.text(x, y, text, ha='center', va='center',
                fontsize=10, color=color,
                bbox=dict(boxstyle='round,pad=0.35', fc='white',
                          ec='#bbb', lw=0.8))

    def step_circle(x, y, num, color='#1F4E78'):
        c = plt.Circle((x, y), 2.0, color=color, zorder=5)
        ax.add_patch(c)
        ax.text(x, y, str(num), ha='center', va='center',
                color='white', fontsize=12, fontweight='bold', zorder=6)

    # Title
    ax.text(60, 76, 'How the System Works',
            ha='center', va='center', fontsize=18,
            fontweight='bold', color='#1F4E78')

    # ── Boxes ─────────────────────────────────────────────────
    # Top row: YOU  BROWSER  CAMERA
    box(4,   58, 20, 10, 'YOU\n(User at PC)',         C_USER)
    box(45,  58, 30, 10, 'BROWSER\n(Dashboard page)', C_BROWSER)
    box(96,  58, 20, 10, 'CAMERA\n(on the wall)',     C_CAM)

    # Middle: MEDIA MTX
    box(45,  33, 30, 10, 'MEDIA MTX\n(Video Converter)', C_MTX)

    # Bottom: AI ENGINE
    box(45,   8, 30, 10, 'AI ENGINE\n(Face • Person • Fire)', C_AI)

    # ── Arrows + labels + step numbers ────────────────────────

    # 1: YOU → BROWSER
    arrow(24, 63, 45, 63)
    label(34, 67, 'Opens dashboard')
    step_circle(34, 60, 1)

    # 2: BROWSER → CAMERA (add URL)
    arrow(75, 63, 96, 63)
    label(86, 67, 'Paste camera URL')
    step_circle(86, 60, 2)

    # 3: CAMERA → MEDIA MTX
    arrow(106, 58, 75, 43)
    label(96, 51, 'Camera video (RTSP)')
    step_circle(88, 55, 3)

    # 4: MEDIA MTX → BROWSER (up)
    arrow(55, 43, 55, 58)
    label(38, 50, 'Browser-friendly video')
    step_circle(58, 50, 4)

    # 5: MEDIA MTX → AI ENGINE (down)
    arrow(65, 33, 65, 18)
    label(84, 25, 'Same video for AI')
    step_circle(68, 25, 5)

    # 6: AI ENGINE → BROWSER (alerts, curved via left)
    arrow(45, 15, 14, 58, color='#C73E1D')
    label(24, 32, 'Live alerts & boxes', color='#C73E1D')
    step_circle(31, 40, 6, color='#C73E1D')

    # ── Legend at bottom ──────────────────────────────────────
    ax.text(60, 2,
            'Steps 1-2: you set it up.   Steps 3-4: video flows to your screen.   '
            'Steps 5-6: AI adds smart alerts on top.',
            ha='center', va='center', fontsize=10, style='italic', color='#333')

    plt.tight_layout()
    plt.savefig(path, dpi=170, bbox_inches='tight', facecolor=C_BG)
    plt.close()
    print(f'Diagram saved: {path}')


# ── DOCX helpers ───────────────────────────────────────────────
def shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def style_header_row(row):
    for cell in row.cells:
        shade(cell, '1F4E78')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(10)


def add_heading(doc, text, level, color=BLUE):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h


def add_para(doc, text, bold=False, italic=False, size=10):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.bold = True
        r.font.size = Pt(10)
        r = p.add_run(' — ' + text)
        r.font.size = Pt(10)
    else:
        r = p.add_run(text)
        r.font.size = Pt(10)
    return p


def add_table(doc, headers, rows, col_widths=None, small=False):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    style_header_row(t.rows[0])
    body_size = Pt(9) if small else Pt(10)
    for r_idx, row_data in enumerate(rows):
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            if r_idx % 2 == 1:
                shade(row.cells[i], 'F2F6FA')
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = body_size
                    if i == 0:
                        run.font.bold = True
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    return t


def page_break(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)


# ── Main ───────────────────────────────────────────────────────
def main():
    diagram_path = 'system_flow.png'
    make_diagram(diagram_path)

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # ── Title ──
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run('Cloud Cost & System Flow')
    r.font.size = Pt(22); r.font.bold = True; r.font.color.rgb = BLUE

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = s.add_run('Comfortable Cloud Setup Pricing + How the Platform Works')
    r.font.size = Pt(12); r.font.italic = True; r.font.color.rgb = GREY

    # ══════════════════════════════════════════════════════════
    # PART 1 — Comfortable Cloud Cost
    # ══════════════════════════════════════════════════════════
    add_heading(doc, 'Part 1: Comfortable Cloud Setup — Full Cost', level=1)
    add_para(doc,
        'The "Comfortable" tier is what we recommend if you want no compromises: '
        '10 cameras with every AI feature on, room to grow to 20+ cameras, and reliable '
        '24/7 uptime. Numbers below are for AWS in a US region (Azure and Google are '
        'within 5-10%).'
    )

    add_heading(doc, 'Monthly Cost Breakdown', level=2)
    add_table(doc,
        ['Item', 'What It Is', 'On-Demand', '1-Year Reserved'],
        [
            ['GPU Server',
             'AWS g5.xlarge — 4 vCPU, 16 GB RAM, NVIDIA A10G (24 GB)',
             '$730',
             '$460'],
            ['Extra Storage',
             '500 GB NVMe SSD attached to the server',
             '$50',
             '$50'],
            ['Managed Database',
             'AWS RDS MySQL, small instance with automated backups',
             '$80',
             '$55'],
            ['Snapshot Storage',
             'S3 bucket for photos — ~50 GB rolling',
             '$5',
             '$5'],
            ['Backup Storage',
             'Daily database + config backups, kept 30 days',
             '$10',
             '$10'],
            ['Bandwidth (outbound)',
             '1 person watching feeds full-time from outside the office',
             '$250-300',
             '$250-300'],
            ['Monitoring',
             'CloudWatch alerts, uptime pings',
             '$15',
             '$15'],
            ['Static IP + DNS',
             'Fixed address for cameras to push to',
             '$5',
             '$5'],
            ['TOTAL',
             'What lands on your card each month',
             '~$1,145',
             '~$850'],
        ],
        col_widths=[Inches(1.5), Inches(3.4), Inches(1.1), Inches(1.3)],
        small=True
    )
    add_para(doc,
        'Reserved pricing (paying for a full year up front or committing monthly) '
        'gives about a 35% discount. If you\'re certain you\'ll run this for 12+ months, '
        'take that path.', italic=True, size=9
    )

    add_heading(doc, 'What Drives the Cost Up or Down', level=2)
    add_bullet(doc, 'How many people watch feeds from outside your LAN. Each extra continuous '
                    'viewer adds $150-300 per month in bandwidth.',
               bold_prefix='Bandwidth')
    add_bullet(doc, 'On-demand is easiest to start with. Commit to 1-3 years once you\'re '
                    'confident and save 35-60%.',
               bold_prefix='Reserved vs on-demand')
    add_bullet(doc, 'Cheaper providers like Runpod, Vast.ai, or Lambda Labs can run the same '
                    'GPU at 40-60% less, but with less enterprise polish.',
               bold_prefix='Alternative clouds')
    add_bullet(doc, 'Every 10 additional cameras roughly adds another $50-100 in bandwidth '
                    'and no other cost until you cross ~30 total.',
               bold_prefix='Scaling up')

    add_heading(doc, 'One-Year Total', level=2)
    add_para(doc,
        'Roughly $10,000-$14,000 per year for the Comfortable tier. For comparison, '
        'a comparable Milestone XProtect deployment for 10 cameras with face recognition '
        '(via BriefCam) typically costs $8,000-$15,000 upfront in licenses PLUS annual '
        'support fees PLUS hardware. Our platform has no license fees — you only pay for '
        'the cloud infrastructure.', bold=True
    )

    # ══════════════════════════════════════════════════════════
    # PART 2 — System diagram + walkthrough
    # ══════════════════════════════════════════════════════════
    page_break(doc)

    add_heading(doc, 'Part 2: How the System Works', level=1)
    add_para(doc,
        'A simple walk-through of what happens when you use the dashboard. '
        'No technical jargon — just what each part does and how they connect.'
    )

    doc.add_picture(diagram_path, width=Inches(7.0))
    caption = doc.paragraphs[-1]
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_heading(doc, 'Step by Step', level=2)

    steps = [
        ('1.  You open the dashboard',
         'You visit our website in a normal browser like Chrome or Edge. '
         'You log in with a username and password.'),
        ('2.  You add a camera',
         'You paste your camera\'s address (called an RTSP URL) into the dashboard '
         'and give it a name like "Front Door" or "Warehouse".'),
        ('3.  The camera streams to Media MTX',
         'Media MTX is a small helper that receives the raw camera video. '
         'Think of it as the "video mail room" — it takes video in one format '
         'and gets it ready for the browser.'),
        ('4.  Media MTX sends the video to your browser',
         'The video is converted into a format any browser can play, and it '
         'shows up live in your dashboard tile — usually within 5-10 seconds '
         'of when it happened in real life.'),
        ('5.  The AI Engine watches the same video',
         'While you watch the video, our AI engine also watches it in the '
         'background — looking for faces, people, and fire.'),
        ('6.  Alerts appear on your dashboard',
         'When the AI sees something interesting — a known face, an unknown '
         'person, or smoke — it instantly draws a box on the video and adds '
         'the event to your alert list on the side of the dashboard.'),
    ]
    for title, body in steps:
        p = doc.add_paragraph()
        r = p.add_run(title)
        r.font.bold = True; r.font.size = Pt(11); r.font.color.rgb = BLUE
        p2 = doc.add_paragraph()
        r = p2.add_run(body)
        r.font.size = Pt(10)

    add_heading(doc, 'The Simple Version', level=2)
    add_para(doc,
        'Cameras send video → Media MTX makes it browser-ready → your browser plays it → '
        'the AI Engine watches at the same time → alerts pop up on your screen. All of '
        'this happens live, on one server, whether that server is in your office or in '
        'the cloud.', italic=True
    )

    out = 'CLOUD_AND_FLOW.docx'
    doc.save(out)
    print(f'Saved {out}')


if __name__ == '__main__':
    main()
