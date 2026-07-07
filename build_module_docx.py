"""Build MODULE_COMPARISON.docx — module-wise, simple-language comparison."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def shade(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def style_header_row(row):
    for cell in row.cells:
        shade(cell, '1F4E78')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_table(doc, headers, rows, col_widths=None, zebra=True):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = h
    style_header_row(t.rows[0])
    for r_idx, row_data in enumerate(rows):
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            if zebra and r_idx % 2 == 1:
                shade(row.cells[i], 'F2F6FA')
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph()
    return t


def main():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('Module-by-Module Comparison')
    r.font.size = Pt(24)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Our AI-CCTV Dashboard  vs.  Milestone XProtect')
    r.font.size = Pt(13)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()
    add_para(doc,
        'A simple, side-by-side look at each part of the system. '
        'No technical jargon — just what each module does, and how the two products handle it.'
    )

    # ── Big table: module by module ──
    headers = ['Module', 'Our Project', 'Milestone XProtect']
    rows = [
        ['Login & Users',
         'Username and password. Each user only sees their own cameras and data.',
         'Big company login with single sign-on, Active Directory, and detailed roles.'],

        ['Camera Setup',
         'Paste an RTSP link and a name. Camera is added in a few seconds.',
         'Wizard with auto-discovery, drivers for thousands of camera brands.'],

        ['Live View',
         'See all your cameras in a browser grid. Refreshes a few frames a second.',
         'Smooth live view in a desktop or mobile app, hardware-accelerated.'],

        ['Recording',
         'No 24/7 recording. We only save photos when something important happens.',
         'Records all cameras non-stop for days, weeks, or months.'],

        ['Face Recognition',
         'Built in. Upload photos of a person, the system learns their face right away.',
         'Needs a paid add-on (BriefCam). Extra setup and license required.'],

        ['Person Detection',
         'Built in. Draws a box around every person it sees.',
         'Available through add-ons.'],

        ['Fire & Smoke',
         'Built in. Detects fire with a smart AI model, or falls back to color detection.',
         'Available through 3rd-party plugins.'],

        ['Alerts & Incidents',
         'Every important event is logged with time, camera, and a snapshot.',
         'Full alarm manager with maps, instructions, and follow-up actions.'],

        ['Snapshots / Evidence',
         'Saves a photo when a face is recognized or fire is seen. Marked with boxes.',
         'Locked, encrypted, and digitally signed evidence ready for court.'],

        ['Map View',
         'Drag and drop your cameras on a world map (Leaflet).',
         'Smart Map with multiple floors, indoor plans, and large GIS layers.'],

        ['Analytics & Reports',
         'Heatmaps, daily counts, top people, top cameras. Click a person to see their path.',
         'Centralized search across millions of events with deep filtering.'],

        ['Person Profiles',
         'Tag each person as VIP, Staff, Visitor, Threat, or Standard. Add notes.',
         'No direct equivalent — handled via custom integrations.'],

        ['Mobile Access',
         'Open the dashboard URL on a phone browser. No native app.',
         'Native iOS and Android apps with push notifications.'],

        ['Multi-Site',
         'One server. Best for a single building or shop.',
         'Connects many buildings or cities into one screen.'],

        ['Security',
         'Login token (JWT), per-user data separation.',
         'Bank-grade encryption, signed clips, role-based permissions, audit logs.'],

        ['Storage',
         'A folder of snapshot photos. MySQL holds the records.',
         'Dedicated recording servers, failover, hot/cold storage tiers.'],

        ['Audio',
         'No audio support.',
         'Two-way audio: listen in, talk back through cameras.'],

        ['License Plates',
         'Not supported.',
         'Built-in License Plate Recognition (LPR).'],

        ['Integrations',
         'Open source — anyone can modify the code.',
         'Open platform with 1,000+ ready-made third-party integrations.'],

        ['Setup / Install',
         'Run start.bat. Done in under a minute.',
         'Enterprise install with planning, sizing, and training.'],

        ['Cost',
         'Free.',
         'Paid license per camera, plus add-ons.'],

        ['Support',
         'Self-help. You own the code.',
         '24/7 paid support contracts with guaranteed response times.'],
    ]

    add_table(doc, headers, rows,
              col_widths=[Inches(1.5), Inches(2.7), Inches(2.7)])

    # ── Short verdict ──
    add_heading(doc, 'The Short Verdict', level=1)
    add_para(doc,
        'Our project is a smart, lightweight watchman — it spots faces, fire, and people in '
        'real time and tells you instantly. It is free, easy to run, and you control everything.'
    )
    add_para(doc,
        'XProtect is a heavy-duty enterprise system — it records everything, scales to '
        'thousands of cameras, and is built for legal-grade evidence. But you pay for that '
        'power, and the AI features cost extra.'
    )
    add_para(doc,
        'Use ours for small sites that need live AI alerts. Use XProtect for big sites that '
        'need archives, audits, and many users.', italic=True
    )

    out = 'MODULE_COMPARISON.docx'
    doc.save(out)
    print(f'Saved {out}')


if __name__ == '__main__':
    main()
