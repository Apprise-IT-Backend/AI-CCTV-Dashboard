"""Build FEATURE_COMPARISON.docx from the comparison content."""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def style_header_row(row):
    for cell in row.cells:
        set_cell_shading(cell, '1F4E78')
        for p in cell.paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(11)


def add_heading(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11)
    return p


def add_table(doc, headers, rows, col_widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
    style_header_row(t.rows[0])
    for row_data in rows:
        row = t.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = val
            for p in row.cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                row.cells[i].width = w
    doc.add_paragraph()
    return t


def main():
    doc = Document()

    # Default font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ── Title ──
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('Feature Comparison')
    r.font.size = Pt(26)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x78)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run('Our AI-CCTV Dashboard  vs.  Milestone XProtect')
    r.font.size = Pt(14)
    r.font.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    doc.add_paragraph()

    add_para(doc,
        'This document compares our project with Milestone XProtect, one of the biggest '
        'enterprise video management systems (VMS) in the world. The goal is to show, in '
        'plain language, what each system does well and where the gaps are.'
    )

    # ── At a glance ──
    add_heading(doc, 'At a Glance', level=1)
    add_table(doc,
        ['', 'Our Project', 'Milestone XProtect'],
        [
            ['What it is',  'AI-first live dashboard',                'Enterprise recording + management VMS'],
            ['Best at',     'Real-time face / fire / person detection','Long-term recording, evidence, scale'],
            ['Scale',       'One server, a handful of cameras',       'Thousands of cameras, many sites'],
            ['Cost',        'Free, open-source',                       'Paid, per-camera licensing'],
            ['Setup',       'Run start.bat, done',                     'Enterprise deployment + training'],
        ],
        col_widths=[Inches(1.3), Inches(2.6), Inches(2.6)]
    )

    # ── What both systems can do ──
    add_heading(doc, 'What Both Systems Can Do', level=1)
    add_para(doc,
        'These are features where both products cover the same ground (even if the depth differs).'
    )
    add_table(doc,
        ['Feature', 'Our Project', 'XProtect'],
        [
            ['View multiple cameras live in one screen',  'Yes (HLS in browser)',       'Yes (Smart Client)'],
            ['Connect to RTSP IP cameras',                'Yes (via MediaMTX)',         'Yes'],
            ['Detect people in the frame',                'Yes (YOLOv8)',               'Yes (via add-ons)'],
            ['Recognize known faces by name',             'Yes (FaceNet, built-in)',    'Yes (via BriefCam add-on)'],
            ['Detect fire and smoke',                     'Yes (built-in)',             'Yes (via add-ons)'],
            ['Save alert snapshots',                      'Yes (auto-burned boxes)',    'Yes (Evidence Lock)'],
            ['Log incidents to a database',               'Yes (MySQL)',                'Yes'],
            ['Show cameras on a map',                     'Yes (Leaflet)',              'Yes (Smart Map)'],
            ['Multiple users with logins',                'Yes (JWT)',                  'Yes'],
            ['Per-user data separation',                  'Yes',                        'Yes (via roles)'],
            ['Charts and analytics',                      'Yes (heatmaps, breakdowns)', 'Yes (Centralized Search)'],
            ['Health monitoring',                         'Basic (/system-health)',     'Yes (Milestone Care)'],
        ],
        col_widths=[Inches(2.4), Inches(2.1), Inches(2.0)]
    )

    # ── What XProtect has that we don't ──
    add_heading(doc, "What XProtect Has That We Don't", level=1)
    add_para(doc,
        "These are the features we'd need to add to compete in the enterprise space."
    )

    add_heading(doc, 'Recording and Storage', level=2)
    for s in [
        'Continuous 24/7 recording — we only save snapshots when something triggers',
        'Retention policies — automatic deletion of old footage by age or disk usage',
        'Failover recording — backup recorder takes over if the main one fails',
        'Hot/cold storage tiers — move old recordings to slower, cheaper disks',
    ]:
        add_bullet(doc, s)

    add_heading(doc, 'Scale and Deployment', level=2)
    for s in [
        'Multi-site federation — connect dozens of buildings into one view (XProtect Interconnect)',
        'Cloud and hybrid hosting — runs on AWS, Azure, Google Cloud',
        'Dedicated hardware appliances — Husky servers pre-tuned for the software',
        'Mobile gateway in DMZ — secure remote access architecture',
    ]:
        add_bullet(doc, s)

    add_heading(doc, 'Security and Evidence', level=2)
    for s in [
        'AES-256 encryption of recorded video at rest',
        'SHA-2 digital signing so exported clips can be proven untampered',
        'Evidence Lock to prevent accidental or deliberate deletion',
        'Chain-of-custody documentation for use in court',
        'Single sign-on (SSO) with OAuth2 / OpenID Connect / Active Directory',
        'Role-based permissions with granular per-camera, per-feature controls',
    ]:
        add_bullet(doc, s)

    add_heading(doc, 'Client Apps and Viewing', level=2)
    for s in [
        'Native desktop client (Smart Client) with hardware acceleration',
        'Native mobile apps for iOS and Android',
        'Video wall support for control rooms',
        'Time-synchronized playback across many cameras at once',
        'Audio support — two-way audio, recording, playback',
    ]:
        add_bullet(doc, s)

    add_heading(doc, 'Analytics and Integrations', level=2)
    for s in [
        'License Plate Recognition (LPR) built-in',
        '1,000+ third-party integrations (access control, intercoms, sensors)',
        'Open SDK and APIs for custom apps',
        'Access control integration (door entry, intercoms)',
    ]:
        add_bullet(doc, s)

    add_heading(doc, 'Enterprise Operations', level=2)
    for s in [
        '24/7 support contracts with guaranteed response times',
        'Centralized management of cameras across all sites',
    ]:
        add_bullet(doc, s)

    # ── What we have that XProtect doesn't ──
    add_heading(doc, "What We Have That XProtect Doesn't (Out of the Box)", level=1)
    add_para(doc,
        'These are real wins for our project — XProtect usually needs paid add-ons for these.'
    )
    add_table(doc,
        ['Feature', 'Why It Matters'],
        [
            ['Built-in face recognition (FaceNet)',
             'No paid add-on, no separate analytics server. Just upload photos and it works.'],
            ['Hot-reload of enrolled faces',
             'Add a new face, the AI picks it up within seconds — no restart.'],
            ['Built-in fire/smoke detection',
             'Custom YOLO model with a smart HSV fallback when no model is installed.'],
            ['Per-person categorization',
             'Tag people as standard / staff / vip / visitor / threat with notes.'],
            ['Movement-path drill-down',
             "Click a recognized person → see their path across cameras on a map."],
            ['Per-user feature toggles',
             'Each user picks which detections to log (face / person / fire).'],
            ['Auto-throttled incident snapshots',
             "Smart per-type throttling so fires don't get drowned out by face matches."],
            ['Single-command startup',
             'start.bat launches the whole stack.'],
            ['No licensing fees',
             'Open source, runs on any modest Windows or Linux box.'],
        ],
        col_widths=[Inches(2.3), Inches(4.2)]
    )

    # ── When to use which ──
    add_heading(doc, 'When to Use Which', level=1)

    add_para(doc, 'Pick our project if you want:', bold=True)
    for s in [
        'A working AI dashboard today, free of charge',
        'Real-time face and fire alerts as the main feature',
        'A small site with under ~20 cameras',
        'Full control over the code and data',
    ]:
        add_bullet(doc, s)

    add_para(doc, 'Pick XProtect if you need:', bold=True)
    for s in [
        '24/7 recorded video for weeks or months',
        'Many sites or many hundreds of cameras',
        'Court-admissible evidence with chain of custody',
        'Enterprise SSO, audit logs, support contracts',
        'A mobile app for guards on the go',
    ]:
        add_bullet(doc, s)

    # ── Honest summary ──
    add_heading(doc, 'Honest Summary', level=1)
    add_para(doc,
        'XProtect is a recording-first platform that the industry has built on for 20+ years. '
        'It records everything, then lets you search and analyze later. Analytics like face '
        'recognition are bolted on through paid partners.'
    )
    add_para(doc,
        'Our project is the opposite — analytics-first. It watches cameras live, reacts to '
        "events instantly, and only saves what matters. There's no archive, no federation, no "
        'signed evidence pipeline. But the AI features that XProtect charges extra for are '
        'built right in and free to run.'
    )
    add_para(doc,
        'For a small business, school, or single building that wants smart alerts without '
        'paying for an enterprise VMS, our project covers the day-to-day needs. For anything '
        'with legal, multi-site, or 24/7 recording requirements, XProtect is the safer choice.'
    )

    out = 'FEATURE_COMPARISON.docx'
    doc.save(out)
    print(f'Saved {out}')


if __name__ == '__main__':
    main()
